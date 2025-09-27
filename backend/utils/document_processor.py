import io
import os
import tempfile
from typing import Dict, Any, Optional, List, Tuple
from PIL import Image
import pytesseract
import PyPDF2
from docx import Document
import openpyxl
from openpyxl.utils import get_column_letter
import logging

# Import the enhanced OCR service
from .enhanced_ocr_service import EnhancedOCRService, OCRResult, OCRQuality

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Service for processing uploaded documents and extracting content"""
    
    # Supported file types and their MIME types
    SUPPORTED_TYPES = {
        '.pdf': 'application/pdf',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        '.png': 'image/png',
        '.jpg': 'image/jpeg',
        '.jpeg': 'image/jpeg'
    }
    
    # Maximum file size (20MB)
    MAX_FILE_SIZE = 20 * 1024 * 1024
    
    # Initialize enhanced OCR service
    _enhanced_ocr_service = None
    
    @classmethod
    def get_enhanced_ocr_service(cls) -> EnhancedOCRService:
        """Get or create the enhanced OCR service instance"""
        if cls._enhanced_ocr_service is None:
            cls._enhanced_ocr_service = EnhancedOCRService()
        return cls._enhanced_ocr_service
    
    @classmethod
    def is_supported_file(cls, filename: str, content_type: str) -> bool:
        """Check if file type is supported"""
        file_ext = os.path.splitext(filename.lower())[1]
        return file_ext in cls.SUPPORTED_TYPES
    
    @classmethod
    def get_file_extension(cls, filename: str) -> str:
        """Get file extension from filename"""
        return os.path.splitext(filename.lower())[1]
    
    @classmethod
    def validate_file_size(cls, file_size: int) -> bool:
        """Validate file size is within limits"""
        return file_size <= cls.MAX_FILE_SIZE
    
    @classmethod
    async def process_document(cls, file_content: bytes, filename: str, content_type: str) -> Dict[str, Any]:
        """
        Process uploaded document and extract content
        
        Args:
            file_content: Raw file bytes
            filename: Original filename
            content_type: MIME type
            
        Returns:
            Dictionary containing extracted content and metadata
        """
        try:
            file_ext = cls.get_file_extension(filename)
            
            if file_ext == '.pdf':
                return await cls._process_pdf(file_content)
            elif file_ext == '.docx':
                return await cls._process_docx(file_content)
            elif file_ext == '.xlsx':
                return await cls._process_xlsx(file_content)
            elif file_ext in ['.png', '.jpg', '.jpeg']:
                return await cls._process_image(file_content, filename)
            else:
                raise ValueError(f"Unsupported file type: {file_ext}")
                
        except Exception as e:
            logger.error(f"Error processing document {filename}: {str(e)}")
            return {
                'extracted_text': None,
                'extracted_data': None,
                'page_count': None,
                'processing_status': 'failed',
                'processing_error': str(e),
                'key_insights': [],
                'suggested_nodes': []
            }
    
    @classmethod
    async def _process_pdf(cls, file_content: bytes) -> Dict[str, Any]:
        """Process PDF file and extract text"""
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            extracted_text = ""
            page_count = len(pdf_reader.pages)
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        extracted_text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                except Exception as e:
                    logger.warning(f"Error extracting text from PDF page {page_num + 1}: {str(e)}")
                    continue
            
            # Generate insights and suggested nodes
            key_insights = cls._generate_key_insights(extracted_text)
            suggested_nodes = cls._generate_suggested_nodes(extracted_text, page_count)
            
            return {
                'extracted_text': extracted_text.strip(),
                'extracted_data': {'page_count': page_count},
                'page_count': page_count,
                'processing_status': 'completed',
                'processing_error': None,
                'key_insights': key_insights,
                'suggested_nodes': suggested_nodes
            }
            
        except Exception as e:
            raise Exception(f"PDF processing failed: {str(e)}")
    
    @classmethod
    async def _process_docx(cls, file_content: bytes) -> Dict[str, Any]:
        """Process DOCX file and extract text"""
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            
            extracted_text = ""
            paragraph_count = 0
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    extracted_text += paragraph.text + "\n"
                    paragraph_count += 1
            
            # Extract tables if any
            table_data = []
            for table in doc.tables:
                table_rows = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_rows.append(row_data)
                table_data.append(table_rows)
            
            # Generate insights and suggested nodes
            key_insights = cls._generate_key_insights(extracted_text)
            suggested_nodes = cls._generate_suggested_nodes(extracted_text, 1)
            
            return {
                'extracted_text': extracted_text.strip(),
                'extracted_data': {
                    'paragraph_count': paragraph_count,
                    'tables': table_data
                },
                'page_count': 1,  # DOCX doesn't have clear page breaks
                'processing_status': 'completed',
                'processing_error': None,
                'key_insights': key_insights,
                'suggested_nodes': suggested_nodes
            }
            
        except Exception as e:
            raise Exception(f"DOCX processing failed: {str(e)}")
    
    @classmethod
    async def _process_xlsx(cls, file_content: bytes) -> Dict[str, Any]:
        """Process XLSX file and extract data"""
        try:
            xlsx_file = io.BytesIO(file_content)
            workbook = openpyxl.load_workbook(xlsx_file, data_only=True)
            
            extracted_text = ""
            extracted_data = {'sheets': {}}
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_data = []
                sheet_text = f"\n--- Sheet: {sheet_name} ---\n"
                
                # Get all rows with data
                for row in sheet.iter_rows(values_only=True):
                    if any(cell is not None for cell in row):
                        row_data = [str(cell) if cell is not None else "" for cell in row]
                        sheet_data.append(row_data)
                        sheet_text += " | ".join(row_data) + "\n"
                
                extracted_data['sheets'][sheet_name] = sheet_data
                extracted_text += sheet_text
            
            # Generate insights and suggested nodes
            key_insights = cls._generate_key_insights(extracted_text)
            suggested_nodes = cls._generate_suggested_nodes(extracted_text, 1)
            
            return {
                'extracted_text': extracted_text.strip(),
                'extracted_data': extracted_data,
                'page_count': len(workbook.sheetnames),
                'processing_status': 'completed',
                'processing_error': None,
                'key_insights': key_insights,
                'suggested_nodes': suggested_nodes
            }
            
        except Exception as e:
            raise Exception(f"XLSX processing failed: {str(e)}")
    
    @classmethod
    async def _process_image(cls, file_content: bytes, filename: str = "image") -> Dict[str, Any]:
        """Process image file using enhanced OCR service with full content extraction"""
        try:
            # Use enhanced OCR service for comprehensive image processing
            enhanced_ocr = cls.get_enhanced_ocr_service()
            ocr_result = await enhanced_ocr.process_image(file_content, filename)
            
            # Convert OCR result to document processor format
            extracted_text = cls._format_enhanced_ocr_text(ocr_result)
            extracted_data = cls._format_enhanced_ocr_data(ocr_result)
            key_insights = cls._generate_enhanced_insights(ocr_result)
            suggested_nodes = cls._generate_enhanced_nodes(ocr_result, filename)
            
            return {
                'extracted_text': extracted_text,
                'extracted_data': extracted_data,
                'page_count': 1,
                'processing_status': 'completed' if not ocr_result.error_message else 'completed_with_warnings',
                'processing_error': ocr_result.error_message,
                'key_insights': key_insights,
                'suggested_nodes': suggested_nodes,
                'ocr_metadata': {
                    'quality_assessment': ocr_result.quality_assessment.value,
                    'language_detected': ocr_result.language_detected,
                    'confidence_score': ocr_result.confidence_score,
                    'processing_time': ocr_result.processing_time,
                    'regions_detected': len(ocr_result.regions),
                    'tables_detected': len(ocr_result.tables),
                    'chart_elements_detected': len(ocr_result.chart_elements)
                }
            }
            
        except Exception as e:
            logger.error(f"Enhanced image processing failed: {str(e)}")
            # Fallback to basic image processing
            return await cls._process_image_fallback(file_content, filename)
    
    @classmethod
    async def _process_image_fallback(cls, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Fallback image processing when enhanced OCR fails"""
        try:
            image = Image.open(io.BytesIO(file_content))
            
            # Get basic image metadata
            image_data = {
                'width': image.width,
                'height': image.height,
                'format': image.format,
                'mode': image.mode,
                'size_bytes': len(file_content),
                'aspect_ratio': round(image.width / image.height, 2)
            }
            
            # Try basic OCR
            extracted_text = ""
            try:
                extracted_text = pytesseract.image_to_string(image)
                logger.info("Basic OCR text extraction successful")
            except Exception as ocr_error:
                logger.warning(f"Basic OCR failed: {str(ocr_error)}")
                
            # Enhanced image description even without OCR
            if len(extracted_text.strip()) < 10:
                size_mb = round(len(file_content) / (1024 * 1024), 2)
                resolution_category = "high" if image.width > 1920 or image.height > 1080 else "standard"
                
                extracted_text = f"""Image Analysis (Basic):
- Format: {image.format} image file
- Dimensions: {image.width} x {image.height} pixels ({resolution_category} resolution)
- File size: {size_mb} MB
- Aspect ratio: {image_data['aspect_ratio']}:1
- Color mode: {image.mode}

This appears to be a {resolution_category} resolution {image.format.lower()} image. Enhanced OCR processing was not available, so only basic metadata analysis was performed."""
            
            # Generate basic insights
            key_insights = cls._generate_image_insights(extracted_text, image_data)
            suggested_nodes = cls._generate_suggested_nodes(extracted_text, 1)
            
            return {
                'extracted_text': extracted_text.strip(),
                'extracted_data': image_data,
                'page_count': 1,
                'processing_status': 'completed',
                'processing_error': None,
                'key_insights': key_insights,
                'suggested_nodes': suggested_nodes
            }
            
        except Exception as e:
            logger.error(f"Fallback image processing failed: {str(e)}")
            raise Exception(f"Image processing failed: {str(e)}")
    
    @classmethod
    def _format_enhanced_ocr_text(cls, ocr_result: OCRResult) -> str:
        """Format enhanced OCR results into readable text"""
        if not ocr_result.original_text and not ocr_result.regions:
            return "No text content detected in image."
        
        formatted_text = ""
        
        # Add original text if available
        if ocr_result.original_text:
            formatted_text += f"=== EXTRACTED TEXT ===\n{ocr_result.original_text}\n\n"
        
        # Add table data if detected
        if ocr_result.tables:
            formatted_text += "=== DETECTED TABLES ===\n"
            for i, table in enumerate(ocr_result.tables, 1):
                formatted_text += f"Table {i} ({table.rows} rows x {table.cols} columns):\n"
                # Group cells by row
                rows = {}
                for cell in table.cells:
                    if cell.row not in rows:
                        rows[cell.row] = {}
                    rows[cell.row][cell.col] = cell.text
                
                # Format table
                for row_idx in sorted(rows.keys()):
                    row_data = rows[row_idx]
                    row_text = " | ".join(row_data.get(col, "") for col in range(table.cols))
                    formatted_text += f"{row_text}\n"
                formatted_text += "\n"
        
        # Add chart elements if detected
        if ocr_result.chart_elements:
            formatted_text += "=== CHART/GRAPH ELEMENTS ===\n"
            for element in ocr_result.chart_elements:
                formatted_text += f"{element.element_type}: {element.text}\n"
            formatted_text += "\n"
        
        # Add region information with spatial references
        if ocr_result.regions:
            formatted_text += "=== SPATIAL TEXT REGIONS ===\n"
            for region in ocr_result.regions:
                region_ref = region.metadata.get('region_reference', 'unknown')
                formatted_text += f"[{region_ref}] {region.text}\n"
            formatted_text += "\n"
        
        # Add metadata summary
        formatted_text += f"=== OCR METADATA ===\n"
        formatted_text += f"Quality: {ocr_result.quality_assessment.value}\n"
        formatted_text += f"Language: {ocr_result.language_detected}\n"
        formatted_text += f"Confidence: {ocr_result.confidence_score:.1f}%\n"
        formatted_text += f"Processing time: {ocr_result.processing_time:.2f}s\n"
        
        return formatted_text.strip()
    
    @classmethod
    def _format_enhanced_ocr_data(cls, ocr_result: OCRResult) -> Dict[str, Any]:
        """Format enhanced OCR data for storage"""
        return {
            'ocr_quality': ocr_result.quality_assessment.value,
            'language_detected': ocr_result.language_detected,
            'confidence_score': ocr_result.confidence_score,
            'processing_time': ocr_result.processing_time,
            'regions_count': len(ocr_result.regions),
            'tables_count': len(ocr_result.tables),
            'chart_elements_count': len(ocr_result.chart_elements),
            'regions': [
                {
                    'x': region.x, 'y': region.y,
                    'width': region.width, 'height': region.height,
                    'text': region.text, 'confidence': region.confidence,
                    'content_type': region.content_type.value,
                    'region_reference': region.metadata.get('region_reference'),
                    'relative_position': region.metadata.get('relative_position')
                } for region in ocr_result.regions
            ],
            'tables': [
                {
                    'rows': table.rows, 'cols': table.cols,
                    'confidence': table.confidence,
                    'bbox': table.bbox,
                    'cells': [
                        {
                            'row': cell.row, 'col': cell.col,
                            'text': cell.text, 'confidence': cell.confidence,
                            'bbox': cell.bbox
                        } for cell in table.cells
                    ]
                } for table in ocr_result.tables
            ],
            'chart_elements': [
                {
                    'type': element.element_type,
                    'text': element.text,
                    'position': element.position,
                    'confidence': element.confidence,
                    'metadata': element.metadata
                } for element in ocr_result.chart_elements
            ]
        }
    
    @classmethod
    def _generate_enhanced_insights(cls, ocr_result: OCRResult) -> List[str]:
        """Generate insights from enhanced OCR results"""
        insights = []
        
        # Quality-based insights
        if ocr_result.quality_assessment == OCRQuality.EXCELLENT:
            insights.append("Excellent image quality - high confidence OCR results")
        elif ocr_result.quality_assessment == OCRQuality.HIGH:
            insights.append("High image quality - reliable OCR extraction")
        elif ocr_result.quality_assessment == OCRQuality.MEDIUM:
            insights.append("Medium image quality - OCR results may need verification")
        else:
            insights.append("Low image quality - OCR results may be incomplete")
        
        # Content type insights
        if ocr_result.tables:
            insights.append(f"Detected {len(ocr_result.tables)} structured table(s) with data")
        
        if ocr_result.chart_elements:
            chart_types = set(element.element_type for element in ocr_result.chart_elements)
            insights.append(f"Contains chart/graph elements: {', '.join(chart_types)}")
        
        # Language insights
        if ocr_result.language_detected != 'unknown':
            insights.append(f"Primary language detected: {ocr_result.language_detected}")
        
        # Confidence insights
        if ocr_result.confidence_score > 80:
            insights.append("High confidence text extraction")
        elif ocr_result.confidence_score > 60:
            insights.append("Moderate confidence text extraction")
        else:
            insights.append("Low confidence text extraction - manual review recommended")
        
        # Region-based insights
        if len(ocr_result.regions) > 10:
            insights.append("Dense text content with multiple regions")
        elif len(ocr_result.regions) < 3:
            insights.append("Sparse text content - may be diagram or simple image")
        
        return insights[:5]  # Limit to 5 insights
    
    @classmethod
    def _generate_enhanced_nodes(cls, ocr_result: OCRResult, filename: str) -> List[Dict[str, Any]]:
        """Generate suggested nodes from enhanced OCR results"""
        suggested_nodes = []
        
        # Generate nodes from table data
        for i, table in enumerate(ocr_result.tables):
            if table.cells:
                # Use first row as potential headers
                first_row_cells = [cell for cell in table.cells if cell.row == 0]
                if first_row_cells:
                    title = " | ".join(cell.text[:20] for cell in sorted(first_row_cells, key=lambda x: x.col))
                    suggested_nodes.append({
                        'title': f"Table Data: {title}",
                        'description': f"Structured data table extracted from {filename} ({table.rows}x{table.cols})",
                        'type': 'human',
                        'confidence': int(table.confidence),
                        'source_type': 'table',
                        'metadata': {
                            'table_index': i,
                            'rows': table.rows,
                            'cols': table.cols
                        }
                    })
        
        # Generate nodes from chart elements
        chart_titles = [elem for elem in ocr_result.chart_elements if elem.element_type == 'title']
        for title_elem in chart_titles:
            suggested_nodes.append({
                'title': f"Chart: {title_elem.text}",
                'description': f"Chart/graph extracted from {filename}",
                'type': 'human',
                'confidence': int(title_elem.confidence),
                'source_type': 'chart',
                'metadata': {
                    'chart_title': title_elem.text,
                    'position': title_elem.position
                }
            })
        
        # Generate nodes from high-confidence text regions
        high_conf_regions = [r for r in ocr_result.regions if r.confidence > 70 and len(r.text) > 20]
        for region in high_conf_regions[:3]:  # Limit to 3 regions
            suggested_nodes.append({
                'title': region.text[:50] + ("..." if len(region.text) > 50 else ""),
                'description': f"Text region from {filename} ({region.metadata.get('region_reference', 'unknown location')})",
                'type': 'human',
                'confidence': int(region.confidence),
                'source_type': 'text_region',
                'metadata': {
                    'region_reference': region.metadata.get('region_reference'),
                    'relative_position': region.metadata.get('relative_position')
                }
            })
        
        return suggested_nodes[:5]  # Limit to 5 suggestions
    
    @classmethod
    def _generate_image_insights(cls, text: str, image_data: Dict) -> List[str]:
        """Generate insights specifically for image files"""
        insights = []
        
        # Resolution insights
        width, height = image_data['width'], image_data['height']
        if width > 1920 or height > 1080:
            insights.append("High resolution image suitable for detailed analysis")
        elif width < 800 and height < 600:
            insights.append("Lower resolution image, may be an icon or thumbnail")
        else:
            insights.append("Standard resolution image")
        
        # Aspect ratio insights
        aspect_ratio = image_data['aspect_ratio']
        if 1.7 <= aspect_ratio <= 1.8:
            insights.append("Widescreen aspect ratio, likely a presentation slide or screenshot")
        elif 0.9 <= aspect_ratio <= 1.1:
            insights.append("Square or near-square aspect ratio")
        elif aspect_ratio > 2:
            insights.append("Very wide aspect ratio, possibly a banner or panoramic image")
        
        # File size insights
        size_mb = round(image_data['size_bytes'] / (1024 * 1024), 2)
        if size_mb > 5:
            insights.append("Large file size suggests high detail or uncompressed image")
        elif size_mb < 0.1:
            insights.append("Small file size, likely optimized or simple graphics")
        
        # Format insights
        if image_data['format'] == 'PNG':
            insights.append("PNG format supports transparency and is often used for screenshots or graphics")
        elif image_data['format'] == 'JPEG':
            insights.append("JPEG format commonly used for photographs")
        
        return insights[:5]  # Limit to 5 insights
    
    @classmethod
    def _generate_key_insights(cls, text: str) -> List[str]:
        """Generate key insights from extracted text"""
        if not text or len(text.strip()) < 50:
            return []
        
        insights = []
        
        # Simple keyword-based insights
        text_lower = text.lower()
        
        if any(word in text_lower for word in ['strategy', 'strategic', 'plan', 'planning']):
            insights.append("Document contains strategic planning content")
        
        if any(word in text_lower for word in ['risk', 'risks', 'threat', 'threats']):
            insights.append("Document discusses risks and threats")
        
        if any(word in text_lower for word in ['decision', 'decisions', 'choose', 'option']):
            insights.append("Document involves decision-making processes")
        
        if any(word in text_lower for word in ['data', 'analysis', 'metrics', 'statistics']):
            insights.append("Document contains analytical data")
        
        if any(word in text_lower for word in ['recommendation', 'suggest', 'propose']):
            insights.append("Document includes recommendations")
        
        # Add word count insight
        word_count = len(text.split())
        insights.append(f"Document contains approximately {word_count} words")
        
        return insights[:5]  # Limit to 5 insights
    
    @classmethod
    def _generate_suggested_nodes(cls, text: str, page_count: int) -> List[Dict[str, Any]]:
        """Generate suggested nodes based on document content"""
        if not text or len(text.strip()) < 50:
            return []
        
        suggested_nodes = []
        text_lower = text.lower()
        
        # Extract potential titles from the first few sentences
        sentences = text.split('.')[:5]
        
        for i, sentence in enumerate(sentences):
            sentence = sentence.strip()
            if len(sentence) > 20 and len(sentence) < 100:
                node_type = "human"  # Default type
                
                # Determine node type based on content
                if any(word in sentence.lower() for word in ['risk', 'threat', 'danger']):
                    node_type = "risk"
                elif any(word in sentence.lower() for word in ['decision', 'choose', 'option']):
                    node_type = "decision"
                elif any(word in sentence.lower() for word in ['depend', 'require', 'need']):
                    node_type = "dependency"
                
                suggested_nodes.append({
                    'title': sentence[:80] + "..." if len(sentence) > 80 else sentence,
                    'description': f"Extracted from uploaded document (page {min(i+1, page_count)})",
                    'type': node_type,
                    'confidence': 70,
                    'source_type': 'document'
                })
        
        return suggested_nodes[:3]  # Limit to 3 suggestions